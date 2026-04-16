"""Build the AgeVision MCA-IV final project report (.docx).

Thin renderer: walks `scripts.report_content.build()` and emits python-docx
elements. All copy lives in `report_content.py` so the pdf renderer stays in
sync automatically.

Output: FINAL_REVIEW_REPORT_NEERAJ.docx at project root.
Format: Anna University CDE - Times New Roman, plain B/W, double-spaced.
"""
from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import report_content  # noqa: E402


# --------------------------------------------------------------------------- #
#                              GLOBAL CONSTANTS                               #
# --------------------------------------------------------------------------- #

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
OUT_PATH = os.path.join(ROOT, "FINAL_REVIEW_REPORT_NEERAJ.docx")
LOGO_PATH = os.path.join(ROOT, "age_vision_logo.png")

FONT = "Times New Roman"
BODY_PT = 12
HEADING_PT = 14


# --------------------------------------------------------------------------- #
#                                  HELPERS                                    #
# --------------------------------------------------------------------------- #

TAMIL_FONT = "Nirmala UI"


def _has_tamil(text):
    return any("\u0b80" <= ch <= "\u0bff" for ch in text or "")


def _set_run_font(run, *, size=BODY_PT, bold=False, italic=False, color=(0, 0, 0)):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    tamil = _has_tamil(run.text)
    cs_font = TAMIL_FONT if tamil else FONT
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), cs_font)
    rFonts.set(qn("w:eastAsia"), cs_font if tamil else FONT)


def _set_para_format(p, *, alignment=None, line_spacing=2.0, first_line_indent=None,
                     space_before=0, space_after=0, keep_with_next=False):
    pf = p.paragraph_format
    if alignment is not None:
        p.alignment = alignment
    if line_spacing == 2.0:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif line_spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif line_spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if keep_with_next:
        pf.keep_with_next = True


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), FONT)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_page_setup(section):
    # A4 (210 x 297 mm) with Anna University CDE margins:
    # Top 32mm, Bottom 28mm, Left 38mm, Right 23mm.
    # Header 20mm from top edge (page number position per spec).
    section.page_height = Twips(16838)
    section.page_width = Twips(11906)
    section.top_margin = Twips(1814)
    section.bottom_margin = Twips(1587)
    section.left_margin = Twips(2155)
    section.right_margin = Twips(1304)
    section.header_distance = Twips(1134)
    section.footer_distance = Twips(720)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    _set_para_format(p, line_spacing=1.0)


def _add_field(paragraph, instr_text):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    _set_run_font(run, size=BODY_PT)
    return run


def configure_header_with_pagenum(section, *, fmt="lowerRoman", start=1,
                                  title_pg=False):
    """Put a right-aligned PAGE field in the section header and configure
    the page-number format (lowerRoman for prelims, decimal for main text).
    """
    sectPr = section._sectPr

    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))
    pgNumType.set(qn("w:fmt"), fmt)

    if title_pg:
        if sectPr.find(qn("w:titlePg")) is None:
            sectPr.append(OxmlElement("w:titlePg"))
        # Ensure a blank first-page header exists so the title page shows
        # no page number (still counted as page i per the spec).
        first_header = section.first_page_header
        for p in list(first_header.paragraphs):
            for child in list(p._p):
                p._p.remove(child)
        fp = first_header.paragraphs[0]
        _set_para_format(fp, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         line_spacing=1.0)

    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        for child in list(p._p):
            p._p.remove(child)
    p = header.paragraphs[0]
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.0)
    _add_field(p, "PAGE   \\* MERGEFORMAT")

    # Blank the footer so nothing collides with the body text bottom margin.
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        for child in list(p._p):
            p._p.remove(child)


def start_main_text_section(doc):
    """Insert a next-page section break and configure the new section
    to restart page numbering at 1 in Arabic numerals. Returns the new
    section so callers can further adjust it if needed.
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_setup(new_section)
    configure_header_with_pagenum(new_section, fmt="decimal", start=1)
    return new_section


def add_centered_heading(doc, text, *, size=HEADING_PT, bold=True, all_caps=True,
                         italic=False, space_before=0, space_after=12,
                         keep_with_next=True):
    p = doc.add_paragraph()
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                     space_before=space_before, space_after=space_after,
                     keep_with_next=keep_with_next)
    run = p.add_run(text.upper() if all_caps else text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_section_heading(doc, number, text, *, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                     space_before=space_before, space_after=space_after,
                     keep_with_next=True)
    if level == 1:
        text_to_show = f"{number} {text.upper()}" if text else number
    else:
        text_to_show = f"{number} {text}" if text else number
    run = p.add_run(text_to_show)
    _set_run_font(run, size=BODY_PT, bold=True)
    return p


def add_body_paragraph(doc, text, *, indent_first_line=True, justify=True):
    p = doc.add_paragraph()
    align = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    _set_para_format(
        p,
        alignment=align,
        line_spacing=1.5,
        first_line_indent=Inches(0.79) if indent_first_line else Pt(0),
    )
    run = p.add_run(text)
    _set_run_font(run, size=BODY_PT)
    return p


def add_centered_paragraph(doc, text, *, bold=False, italic=False, size=BODY_PT,
                           space_before=0, space_after=0):
    p = doc.add_paragraph()
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15,
                     space_before=space_before, space_after=space_after)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_right_paragraph(doc, text, *, bold=False):
    p = doc.add_paragraph()
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.5,
                     space_before=24, space_after=4)
    run = p.add_run(text)
    _set_run_font(run, size=BODY_PT, bold=bold)
    return p


# ----- Lists ---------------------------------------------------------------- #

def _ensure_numbering_part(doc):
    if not hasattr(doc.part, "numbering_part") or doc.part.numbering_part is None:
        from docx.parts.numbering import NumberingPart
        from docx.opc.constants import RELATIONSHIP_TYPE

        np = NumberingPart.new()
        np.partname = "/word/numbering.xml"
        doc.part.relate_to(np, RELATIONSHIP_TYPE.NUMBERING)

    numbering = doc.part.numbering_part.element
    abs_ids = {a.get(qn("w:abstractNumId")) for a in numbering.findall(qn("w:abstractNum"))}
    num_ids = {n.get(qn("w:numId")) for n in numbering.findall(qn("w:num"))}

    bullet_abs_xml = """
<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="10">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/><w:lvlText w:val="\u2022"/><w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
    <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr>
  </w:lvl>
</w:abstractNum>
""".strip()
    decimal_abs_xml = """
<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="11">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>
    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
  </w:lvl>
</w:abstractNum>
""".strip()
    bullet_num_xml = """
<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="100">
  <w:abstractNumId w:val="10"/>
</w:num>
""".strip()
    decimal_num_xml = """
<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="101">
  <w:abstractNumId w:val="11"/>
</w:num>
""".strip()

    from docx.oxml import parse_xml as _parse
    if "10" not in abs_ids:
        numbering.insert(0, _parse(bullet_abs_xml))
    if "11" not in abs_ids:
        numbering.insert(0, _parse(decimal_abs_xml))
    if "100" not in num_ids:
        numbering.append(_parse(bullet_num_xml))
    if "101" not in num_ids:
        numbering.append(_parse(decimal_num_xml))


def _add_list_paragraph(doc, text, num_id):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5)
    run = p.add_run(text)
    _set_run_font(run, size=BODY_PT)
    return p


def add_bullet_list(doc, items):
    _ensure_numbering_part(doc)
    for item in items:
        _add_list_paragraph(doc, item, num_id=100)


def add_numbered_list(doc, items):
    _ensure_numbering_part(doc)
    for item in items:
        _add_list_paragraph(doc, item, num_id=101)


def add_reference_list(doc, items):
    """Render [1] style numbered references with hanging indent."""
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing=1.5, space_after=4)
        run = p.add_run(f"[{i}] {item}")
        _set_run_font(run, size=BODY_PT)


# ----- Tables --------------------------------------------------------------- #

def _set_cell_borders(cell, sz=4, color="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def _clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def _set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=BODY_PT):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_para_format(p, alignment=align, line_spacing=1.15)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_plain_table(doc, headers, rows, caption=None, *, col_widths=None,
                    borderless=False, bold_rows=None,
                    header_align=WD_ALIGN_PARAGRAPH.CENTER):
    bold_rows = set(bold_rows or [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths:
        total = sum(col_widths)
        for i, col in enumerate(table.columns):
            for cell in col.cells:
                cell.width = Twips(int(9026 * col_widths[i] / total))

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=not borderless, align=header_align)
        if borderless:
            _clear_cell_borders(cell)
        else:
            _set_cell_borders(cell)

    for r_i, row in enumerate(rows, start=1):
        row_bold = (r_i - 1) in bold_rows
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            _set_cell_text(cell, str(val), bold=row_bold,
                           align=WD_ALIGN_PARAGRAPH.LEFT)
            if borderless:
                _clear_cell_borders(cell)
            else:
                _set_cell_borders(cell)

    if caption:
        add_centered_paragraph(doc, caption, bold=True, size=BODY_PT,
                               space_before=4, space_after=8)
    else:
        doc.add_paragraph("")
    return table


# ----- Figure placeholder --------------------------------------------------- #

SCREENSHOT_DIR = os.path.join(ROOT, "screenshots")


def add_figure_placeholder(doc, caption, *, image=None, height_in=1.7):
    img_path = None
    if image:
        candidate = os.path.join(SCREENSHOT_DIR, image)
        if os.path.exists(candidate):
            img_path = candidate

    if img_path:
        p = doc.add_paragraph()
        _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.0, space_before=6, space_after=2)
        run = p.add_run()
        try:
            run.add_picture(img_path, width=Inches(4.5))
        except Exception:
            img_path = None

    if not img_path:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.rows[0].cells[0]
        cell.width = Inches(4.5)
        _set_cell_borders(cell, sz=6)
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), str(int(height_in * 1440)))
        trHeight.set(qn("w:hRule"), "exact")
        trPr.append(trHeight)
        cell_p = cell.paragraphs[0]
        _set_para_format(cell_p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_centered_paragraph(doc, caption, bold=True, size=BODY_PT,
                           space_before=4, space_after=8)


def add_logo(doc, logo_path):
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0,
                         space_before=4, space_after=10)
        run = p.add_run()
        try:
            run.add_picture(logo_path, width=Inches(1.1))
            return
        except Exception:
            pass
    add_centered_paragraph(doc, "[University Logo Placeholder]",
                           italic=True, space_after=10)


def add_mono_block(doc, text):
    p = doc.add_paragraph()
    _set_para_format(p, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=4, space_after=6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), "Consolas")


def add_sub_title(doc, text):
    """Appendix title — bold left-aligned 12pt with extra top space."""
    p = doc.add_paragraph()
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                     space_before=14, space_after=8, keep_with_next=True)
    run = p.add_run(text)
    _set_run_font(run, size=BODY_PT, bold=True)


def add_sub_heading(doc, text):
    """Sub-heading inside an appendix — bold, slightly less space."""
    p = doc.add_paragraph()
    _set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                     space_before=8, space_after=4, keep_with_next=True)
    run = p.add_run(text)
    _set_run_font(run, size=BODY_PT, bold=True)


# --------------------------------------------------------------------------- #
#                                 DISPATCH                                    #
# --------------------------------------------------------------------------- #

def render(doc, commands):
    for cmd in commands:
        op = cmd[0]
        args = cmd[1:]

        if op == "cover_para":
            text, opts = args[0], (args[1] if len(args) > 1 else {})
            add_centered_paragraph(
                doc, text,
                bold=opts.get("bold", False),
                italic=opts.get("italic", False),
                size=opts.get("size", BODY_PT),
                space_before=opts.get("space_before", 0),
                space_after=opts.get("space_after", 6),
            )

        elif op == "logo":
            add_logo(doc, args[0] if args else LOGO_PATH)

        elif op == "heading_centered":
            text, opts = args[0], (args[1] if len(args) > 1 else {})
            add_centered_heading(
                doc, text,
                size=opts.get("size", HEADING_PT),
                bold=opts.get("bold", True),
                italic=opts.get("italic", False),
                all_caps=opts.get("all_caps", False),
                space_before=opts.get("space_before", 0),
                space_after=opts.get("space_after", 12),
            )

        elif op == "section_break":
            # Opts unused for now; always transitions prelims -> main text
            # with decimal numbering restarted at 1.
            start_main_text_section(doc)

        elif op == "section":
            number, text, opts = args[0], args[1], (args[2] if len(args) > 2 else {})
            add_section_heading(doc, number, text, level=opts.get("level", 1))

        elif op == "para":
            add_body_paragraph(doc, args[0])

        elif op == "bullet":
            add_bullet_list(doc, args[0])

        elif op == "numbered":
            add_numbered_list(doc, args[0])

        elif op == "table":
            headers, rows, caption = args[0], args[1], args[2]
            opts = args[3] if len(args) > 3 else {}
            add_plain_table(
                doc, headers, rows, caption,
                col_widths=opts.get("col_widths"),
                borderless=opts.get("borderless", False),
                bold_rows=opts.get("bold_rows"),
            )

        elif op == "figure":
            # New form: (caption, opts_with_image). Legacy: (label, caption).
            if len(args) >= 2 and isinstance(args[1], dict):
                caption, opts = args[0], args[1]
            else:
                caption = args[1] if len(args) > 1 else args[0]
                opts = {}
            add_figure_placeholder(
                doc, caption,
                image=opts.get("image"),
                height_in=opts.get("height_in", 3.0),
            )

        elif op == "vspace":
            pt = args[0]
            p = doc.add_paragraph()
            _set_para_format(p, line_spacing=1.0, space_before=pt / 2,
                             space_after=pt / 2)

        elif op == "pagebreak":
            add_page_break(doc)

        elif op == "mono":
            add_mono_block(doc, args[0])

        elif op == "right_text":
            text, opts = args[0], (args[1] if len(args) > 1 else {})
            add_right_paragraph(doc, text, bold=opts.get("bold", False))

        elif op == "references":
            add_reference_list(doc, args[0])

        elif op == "sub_title":
            add_sub_title(doc, args[0])

        elif op == "sub_heading":
            add_sub_heading(doc, args[0])

        else:
            raise ValueError(f"Unknown DSL command: {op}")


# --------------------------------------------------------------------------- #
#                                   MAIN                                      #
# --------------------------------------------------------------------------- #

def main():
    doc = Document()
    set_default_style(doc)
    set_page_setup(doc.sections[0])
    configure_header_with_pagenum(
        doc.sections[0], fmt="lowerRoman", start=1, title_pg=True,
    )

    commands = report_content.build()
    render(doc, commands)

    out_path = OUT_PATH
    try:
        doc.save(out_path)
    except PermissionError:
        out_path = os.path.join(
            ROOT, "FINAL_REVIEW_REPORT_NEERAJ.new.docx"
        )
        doc.save(out_path)
        print(f"WARNING: {OUT_PATH} was locked (likely open in Word).")
        print(f"         Wrote to {out_path} instead.")
    print(f"Saved: {out_path}")

    doc2 = Document(out_path)
    print(f"Validation OK: {len(doc2.paragraphs)} paragraphs, "
          f"{len(doc2.tables)} tables, {len(doc2.sections)} sections")


if __name__ == "__main__":
    main()
